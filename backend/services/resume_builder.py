"""
Resume Builder Service — uses LLM to generate tailored resume + cover letter.
"""
import json
import os
from sqlalchemy.orm import Session
from backend.models.resume import Resume
from backend.models.job import Job
from backend.models.user import User, ProfileSkill, ProfileExperience

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")


def _build_profile_context(user: User, skills: list, experiences: list) -> str:
    exp_text = "\n".join(
        f"- {e.title} at {e.company} ({e.start_date} - {e.end_date or 'present'}): {e.description or ''}"
        for e in experiences
    )
    skill_text = ", ".join(f"{s.skill} ({s.level})" for s in skills if s.level) or \
                 ", ".join(s.skill for s in skills)
    return f"Name: {user.name}\nEmail: {user.email}\nBio: {user.bio}\nSkills: {skill_text}\nExperience:\n{exp_text}"


async def build(job_id: int, user_id: int, db: Session) -> Resume:
    job = db.query(Job).filter(Job.id == job_id).first()
    user = db.query(User).filter(User.id == user_id).first()
    skills = db.query(ProfileSkill).filter(ProfileSkill.user_id == user_id).all()
    experiences = db.query(ProfileExperience).filter(ProfileExperience.user_id == user_id).all()

    profile_ctx = _build_profile_context(user, skills, experiences)
    content = await _call_llm(profile_ctx, job)

    resume = Resume(
        user_id=user_id,
        job_id=job_id,
        version="1.0",
        content_json=json.dumps(content),
    )
    db.add(resume)
    db.commit()
    db.refresh(resume)

    # Export files
    _export_docx(resume, content)
    _export_pdf(resume, content)
    db.commit()

    return resume


async def _call_llm(profile: str, job: Job) -> dict:
    """Call OpenAI or fall back to a stub."""
    OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
    if not OPENAI_API_KEY:
        return _stub_resume(profile, job)

    try:
        from openai import AsyncOpenAI
        client = AsyncOpenAI(api_key=OPENAI_API_KEY)
        prompt = f"""
You are a professional resume writer. Given the candidate profile and job posting below,
generate a tailored resume and cover letter in JSON format with keys:
summary, skills, experience, education, cover_letter.

CANDIDATE PROFILE:
{profile}

JOB POSTING:
Title: {job.title}
Description: {job.description}
Location: {job.location}
"""
        response = await client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"},
        )
        return json.loads(response.choices[0].message.content)
    except Exception as e:
        import traceback
        print(f"LLM error: {e}")
        traceback.print_exc()
        return _stub_resume(profile, job)


def _stub_resume(profile: str, job: Job) -> dict:
    return {
        "summary": f"Tailored resume for {job.title}",
        "skills": [],
        "experience": [],
        "education": [],
        "cover_letter": f"Dear Hiring Manager, I am applying for {job.title}...",
    }


def _export_docx(resume: Resume, content: dict):
    try:
        from docx import Document
        import os
        os.makedirs("data/resumes", exist_ok=True)
        doc = Document()
        doc.add_heading(content.get("summary", "Resume"), 0)
        doc.add_paragraph(content.get("cover_letter", ""))
        path = f"data/resumes/resume_{resume.id}.docx"
        doc.save(path)
        resume.docx_path = path
    except Exception as e:
        print(f"DOCX export error: {e}")


def _export_pdf(resume: Resume, content: dict):
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        import os
        os.makedirs("data/resumes", exist_ok=True)
        path = f"data/resumes/resume_{resume.id}.pdf"
        doc = SimpleDocTemplate(path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        story.append(Paragraph(content.get("summary", "Resume"), styles["Title"]))
        story.append(Spacer(1, 12))
        if content.get("cover_letter"):
            story.append(Paragraph("Cover Letter", styles["Heading2"]))
            story.append(Paragraph(content["cover_letter"].replace("\n", "<br/>"), styles["Normal"]))
        doc.build(story)
        resume.pdf_path = path
    except Exception as e:
        print(f"PDF export error: {e}")
